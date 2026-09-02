"""
test_template_front_matter.py - 文档前置页（首页 / 修订记录 / 目录 / 页码页脚）测试

覆盖:
- 首页: 文件名称、产品名称、文件编号（阶段编号方案）、版本号、编制/审核/批准
- 修订记录: 6 列表格 + 首行 V1.0 数据
- 目录: Word TOC 域 + updateFields 自动更新
- 页码页脚: PAGE / NUMPAGES 域
- 分页: 首页→修订记录→目录→正文
- 文件编号: 阶段序号推导、未归类兜底、env 前缀覆盖
- 环境变量: 编制/审核/批准/版本号 可配置
- 正文: 仍完整解析 Markdown（标题/表格/列表）
"""
import os
import re
import sys
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

SAMPLE_MARKDOWN = """# 设计输入文件

## 1 产品概述

- 适用范围：持续皮下输注胰岛素
- 预期用途：糖尿病患者血糖管理

### 1.1 功能描述

| 参数 | 限值 |
|---|---|
| 输注精度 | ±5% |

## 2 性能要求

正文段落。
"""

REVISION_HEADERS = ["版本号", "修订日期", "修订内容", "编制", "审核", "批准"]


@pytest.fixture(autouse=True)
def clean_env(monkeypatch):
    """清理前置页相关环境变量，保证测试确定性。"""
    for k in ("DOC_NUMBER_PREFIX", "DOC_VERSION", "DOC_MAKER",
              "DOC_REVIEWER", "DOC_APPROVER"):
        monkeypatch.delenv(k, raising=False)


def _build(monkeypatch=None, doc_type="design_input", product_name="贴敷式胰岛素泵A7"):
    """构建并重新打开一份含前置页的 Document。"""
    from app.services.template import TemplateService
    ts = TemplateService()
    doc = ts.load_template(doc_type)
    ts.fill_template(doc, SAMPLE_MARKDOWN, product_name, doc_type)
    return Document(BytesIO(ts.document_to_bytes(doc)))


def _all_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _page_break_count(doc) -> int:
    return doc.element.xml.count('w:type="page"')


def _find_revision_table(doc):
    for t in doc.tables:
        if len(t.columns) == 6:
            header = [t.cell(0, j).text for j in range(6)]
            if header == REVISION_HEADERS:
                return t
    return None


class TestCoverPage:
    def test_cover_has_doc_number_and_version(self):
        doc = _build()
        text = _all_text(doc)
        assert "文件编号" in text
        assert "版本号" in text
        # design_input 属「二、设计输入」，组内第一个 → KF-CGM-2-0001
        assert "KF-CGM-2-0001" in text
        assert "V1.0" in text

    def test_cover_has_signatures(self):
        doc = _build()
        text = _all_text(doc)
        assert "编制" in text
        assert "审核" in text
        assert "批准" in text
        # 默认占位
        assert "编制人" in text
        assert "审核人" in text
        assert "批准人" in text

    def test_cover_has_label_and_product(self):
        doc = _build()
        text = _all_text(doc)
        assert "设计输入文件" in text
        assert "贴敷式胰岛素泵A7" in text


class TestRevisionHistory:
    def test_revision_table_headers_and_row(self):
        doc = _build()
        table = _find_revision_table(doc)
        assert table is not None
        row = [table.cell(1, j).text for j in range(6)]
        assert row[0] == "V1.0"
        assert row[1] == datetime.now().strftime("%Y-%m-%d")
        assert "首次编制发布" in row[2]
        assert row[3:] == ["编制人", "审核人", "批准人"]


class TestToc:
    def test_toc_field_present(self):
        doc = _build()
        assert 'TOC \\o "1-3"' in doc.element.xml

    def test_update_fields_set(self):
        doc = _build()
        assert "updateFields" in doc.settings.element.xml

    def test_toc_placeholder_text(self):
        doc = _build()
        assert any("目录将在打开文档时自动更新" in p.text for p in doc.paragraphs)


class TestPageFooter:
    def test_page_number_fields(self):
        doc = _build()
        footer_xml = doc.sections[0].footer._element.xml
        assert "PAGE" in footer_xml
        assert "NUMPAGES" in footer_xml


class TestPageBreaks:
    def test_front_matter_separated_by_page_breaks(self):
        doc = _build()
        # 首页→修订记录→目录→正文 至少 3 个分页符
        assert _page_break_count(doc) >= 3


class TestDocNumber:
    def test_stage_number_derivation(self):
        from app.services.template import _derive_doc_number
        assert _derive_doc_number("design_input") == "KF-CGM-2-0001"
        assert _derive_doc_number("design_development_plan") == "KF-CGM-1-0001"
        assert _derive_doc_number("design_validation_master_plan").startswith("KF-CGM-5-")
        assert _derive_doc_number("pms_plan").startswith("KF-CGM-9-")

    def test_legacy_category_uses_stage_0(self):
        from app.services.template import _derive_doc_number
        assert _derive_doc_number("sop").startswith("KF-CGM-0-")

    def test_unknown_type_deterministic_fallback(self):
        from app.services.template import _derive_doc_number
        n1 = _derive_doc_number("nonexistent_type_xyz")
        n2 = _derive_doc_number("nonexistent_type_xyz")
        assert n1 == n2
        assert re.fullmatch(r"KF-CGM-0-\d{4}", n1)

    def test_prefix_env_override(self, monkeypatch):
        from app.services.template import _derive_doc_number
        monkeypatch.setenv("DOC_NUMBER_PREFIX", "ABC-INS")
        assert _derive_doc_number("design_input").startswith("ABC-INS-2-")


class TestEnvOverride:
    def test_signature_names_env(self, monkeypatch):
        monkeypatch.setenv("DOC_MAKER", "张工")
        monkeypatch.setenv("DOC_REVIEWER", "李工")
        monkeypatch.setenv("DOC_APPROVER", "王工")
        doc = _build()
        text = _all_text(doc)
        assert "张工" in text
        assert "李工" in text
        assert "王工" in text

    def test_version_env(self, monkeypatch):
        monkeypatch.setenv("DOC_VERSION", "V2.0")
        doc = _build()
        assert "V2.0" in _all_text(doc)


class TestBodyContent:
    def test_markdown_still_parsed(self):
        doc = _build()
        text = _all_text(doc)
        assert "1 产品概述" in text
        assert "2 性能要求" in text
        assert "1.1 功能描述" in text
        assert "输注精度" in text
        assert "适用范围：持续皮下输注胰岛素" in text

    def test_body_does_not_duplicate_cover_title(self):
        """正文不再重复首页的标题/副标题（避免与首页重复）。"""
        doc = _build()
        # 标题样式段（styleId Title）只应出现一次（首页）
        title_styles = [p for p in doc.paragraphs
                        if p.style is not None and p.style.name == "Title"]
        assert len(title_styles) == 1
