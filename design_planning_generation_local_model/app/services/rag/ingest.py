"""
Ingest - 文档摄入脚本

将 develop_documents 目录下的参考文档处理后写入向量库。
增强版：支持PDF、DOC、DOCX、TXT等多种格式，支持从网络下载的文件摄入。
"""

import os
import sys
import json
import shutil
import hashlib
from pathlib import Path
from typing import Optional, List, Tuple

# 添加项目根目录到 path
project_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from app.services.rag.embedder import Embedder
from app.services.rag.vector_store import VectorStore


# 文档类型映射规则：目录 → doc_type
DOC_TYPE_MAPPING = {
    "CH3.2风险管理模板": "risk_management_report",
    "CH5.03包装说明使用说明书": "instruction",
    "产品技术要求模板": "product_spec",
    "软件资料": "software",
    "体系资料": "system_doc",
    "流程图": "flowchart",
    "作业指导书": "sop",
    "设计开发表格": "design_doc",
    "验证": "validation",
    "有效期": "shelf_life",
    "web/sop": "sop",
    "downloads/sop": "sop",
    "downloads/web/sop": "sop",
}

# 每篇文档最多保留的 chunk 数量（防止单篇文档主导检索结果）
MAX_CHUNKS_PER_DOC = 80  # 大幅增加以保留更多详细内容

# chunk 大小配置
CHUNK_SIZE = 1500  # 增大以保留更完整的语义单元和段落
CHUNK_OVERLAP = 200  # 增加重叠以保持上下文连贯性


def get_doc_type_from_path(file_path: str) -> str:
    """根据文件路径判断文档类型"""
    # 首先尝试完整路径匹配
    for dir_name, doc_type in DOC_TYPE_MAPPING.items():
        if dir_name in file_path.replace("\\", "/"):
            return doc_type
    # 其次尝试文件名关键词
    filename = os.path.basename(file_path).lower()
    if any(keyword in filename for keyword in ["sop", "作业指导", "操作规程"]):
        return "sop"
    if any(keyword in filename for keyword in ["风险", "risk"]):
        return "risk_management_report"
    if any(keyword in filename for keyword in ["说明", "manual"]):
        return "instruction"
    if any(keyword in filename for keyword in ["技术要求", "spec"]):
        return "product_spec"
    return "unknown"


def extract_text_from_docx(docx_path: str) -> List[Tuple[str, str]]:
    """
    从 .docx 文件提取段落文本和章节标题

    Args:
        docx_path: .docx 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("请先安装 python-docx: pip install python-docx")

    doc = Document(docx_path)
    paragraphs = []

    current_section = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 判断是否是标题
        if para.style.name.startswith("Heading"):
            current_section = text
        elif len(text) < 60 and not text.endswith("。") and not text.endswith("."):
            # 短文本可能是标题
            if text.isdigit() or len(text) < 5:
                # 太短的跳过
                continue
            current_section = text
        else:
            paragraphs.append((current_section, text))

    # 处理表格
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_text.append(row_text)
        if table_text:
            paragraphs.append((current_section or "表格", "\n".join(table_text)))

    return paragraphs


def _get_tesseract_path() -> str:
    """获取 tesseract 可执行文件路径"""
    import shutil
    # 尝试多个已知路径
    candidates = [
        r"E:\tesseract\tesseract.exe",
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        shutil.which("tesseract"),
    ]
    for path in candidates:
        if path and os.path.isfile(path):
            return path
    return "tesseract"


def extract_text_from_pdf_ocr(
    pdf_path: str,
    dpi: int = 200,
    max_pages: int = 30,
    lang: str = "chi_sim+eng"
) -> List[Tuple[str, str]]:
    """
    使用 OCR 从扫描版 PDF 中提取文本

    Args:
        pdf_path: PDF 文件路径
        dpi: 渲染分辨率（越高越清晰但越慢）
        max_pages: 最大处理页数
        lang: OCR 语言

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io
    except ImportError as e:
        print(f"  [OCR] 缺少依赖: {e}")
        return []

    pytesseract.pytesseract.tesseract_cmd = _get_tesseract_path()

    paragraphs = []
    blank_pages = 0

    try:
        doc = fitz.open(pdf_path)
        pages_to_process = min(doc.page_count, max_pages)

        for page_num in range(pages_to_process):
            try:
                page = doc[page_num]
                pix = page.get_pixmap(dpi=dpi)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                text = pytesseract.image_to_string(img, lang=lang)

                if text and len(text.strip()) > 20:
                    lines = text.strip().split("\n")
                    current_section = f"第{page_num + 1}页"
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        if len(line) < 80 and not line.endswith((".", "。", "）", ")")):
                            current_section = line
                        else:
                            paragraphs.append((current_section, line))
                    blank_pages = 0
                else:
                    blank_pages += 1
                    # 连续5页空白则停止
                    if blank_pages >= 5:
                        break
            except Exception:
                continue

        doc.close()
    except Exception as e:
        print(f"  [OCR] PDF 处理失败: {e}")

    return paragraphs


def extract_text_from_pdf(pdf_path: str, enable_ocr: bool = True) -> List[Tuple[str, str]]:
    """
    从 .pdf 文件提取文本内容

    Args:
        pdf_path: .pdf 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    try:
        import PyPDF2
    except ImportError:
        try:
            import pdfplumber
        except ImportError:
            print(f"  [PDF] 需要安装 PyPDF2 或 pdfplumber: pip install PyPDF2 pdfplumber")
            return []

    paragraphs = []

    try:
        # 尝试使用 pdfplumber（更好）
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    lines = text.split("\n")
                    current_section = f"第{page_num + 1}页"

                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue

                        # 判断是否可能是标题
                        if (len(line) < 80 and
                            (line.startswith(("第", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.",
                                             "一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、",
                                             "1.1", "1.2", "2.1", "2.2", "3.1", "3.2")) or
                             line.isupper() or
                             (len(line) < 40 and not line.endswith("。")))):
                            current_section = line
                        else:
                            paragraphs.append((current_section, line))

            # pdfplumber 提取到文本才返回，否则继续尝试 PyPDF2 和 OCR
            if paragraphs:
                return paragraphs

    except ImportError:
        pass

    # 回退到 PyPDF2
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    lines = text.split("\n")
                    current_section = f"第{page_num + 1}页"

                    for line in lines:
                        line = line.strip()
                        if line:
                            paragraphs.append((current_section, line))
    except Exception as e:
        print(f"  [PDF] 读取失败: {e}")

    # OCR 回退：如果直接提取无文本，尝试 OCR
    if not paragraphs and enable_ocr:
        print(f"  [PDF] 直接提取无文本，尝试 OCR...")
        ocr_paragraphs = extract_text_from_pdf_ocr(pdf_path)
        if ocr_paragraphs:
            print(f"  [OCR] 成功提取 {len(ocr_paragraphs)} 段文本")
        return ocr_paragraphs

    return paragraphs


def extract_text_from_doc(doc_path: str) -> List[Tuple[str, str]]:
    """
    从 .doc (旧版Word) 文件提取文本

    Args:
        doc_path: .doc 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    try:
        import textract
    except ImportError:
        print(f"  [DOC] 需要安装 textract: pip install textract")
        print(f"  [DOC] (Windows 可能还需要安装 AntiWord 或其他工具)")
        return []

    try:
        text = textract.process(doc_path).decode('utf-8')
        lines = text.split("\n")
        paragraphs = []
        current_section = ""

        for line in lines:
            line = line.strip()
            if not line:
                continue

            if len(line) < 60 and not line.endswith("。"):
                current_section = line
            else:
                paragraphs.append((current_section, line))

        return paragraphs
    except Exception as e:
        print(f"  [DOC] 读取失败: {e}")
        return []


def extract_text_from_xlsx(xlsx_path: str) -> List[Tuple[str, str]]:
    """
    从 .xlsx 文件提取文本内容

    Args:
        xlsx_path: .xlsx 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    try:
        import openpyxl
    except ImportError:
        print(f"  [XLSX] 需要安装 openpyxl: pip install openpyxl")
        return []

    paragraphs = []
    try:
        wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_title = f"[工作表: {sheet_name}]"
            rows_text = []
            for row in ws.iter_rows(values_only=True):
                row_values = [str(cell) for cell in row if cell is not None]
                if row_values:
                    row_text = " | ".join(row_values)
                    rows_text.append(row_text)
            if rows_text:
                paragraphs.append((sheet_title, "\n".join(rows_text)))
        wb.close()
    except Exception as e:
        print(f"  [XLSX] 读取失败: {e}")

    return paragraphs


def extract_text_from_md(md_path: str) -> List[Tuple[str, str]]:
    """
    从 .md 文件提取文本内容
    """
    paragraphs = []
    current_section = ""
    encodings = ['utf-8', 'gbk', 'gb18030', 'latin1']

    for encoding in encodings:
        try:
            with open(md_path, 'r', encoding=encoding) as f:
                lines = f.readlines()

            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Markdown 标题作为 section
                if line.startswith("#"):
                    current_section = line.lstrip("#").strip()
                else:
                    paragraphs.append((current_section, line))
            return paragraphs
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"  [MD] 读取失败: {e}")
            return []

    return paragraphs


def extract_text_from_txt(txt_path: str) -> List[Tuple[str, str]]:
    """
    从 .txt 文件提取文本

    Args:
        txt_path: .txt 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    paragraphs = []
    current_section = ""

    encodings = ['utf-8', 'gbk', 'gb18030', 'latin1']

    for encoding in encodings:
        try:
            with open(txt_path, 'r', encoding=encoding) as f:
                lines = f.readlines()

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if len(line) < 60 and not line.endswith("。") and line:
                    current_section = line
                else:
                    paragraphs.append((current_section, line))

            return paragraphs
        except UnicodeDecodeError:
            continue
        except Exception as e:
            print(f"  [TXT] 读取失败: {e}")
            return []

    return paragraphs


def extract_text_from_file(file_path: str) -> List[Tuple[str, str]]:
    """
    根据文件扩展名选择合适的提取方法

    Args:
        file_path: 文件路径

    Returns:
        [(section_title, paragraph_text), ...] 列表
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.xlsx':
        return extract_text_from_xlsx(file_path)
    elif ext == '.md':
        return extract_text_from_md(file_path)
    else:
        print(f"  [SKIP] 不支持的文件格式: {ext}")
        return []


def chunk_text(
    paragraphs: List[Tuple[str, str]],
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP
) -> List[dict]:
    """
    将段落列表分块

    Args:
        paragraphs: [(section_title, text), ...]
        chunk_size: 每块目标字符数
        chunk_overlap: 块之间的重叠字符数

    Returns:
        [{"chunk_id": str, "text": str, "section_title": str, "chunk_index": int}, ...]
    """
    chunks = []
    current_chunk = ""
    current_section = ""
    chunk_index = 0

    for section_title, text in paragraphs:
        if not text.strip():
            continue

        # 如果单段文本超过 chunk_size，分割它
        if len(text) > chunk_size:
            # 先输出当前累积的 chunk
            if current_chunk:
                chunks.append({
                    "chunk_id": f"chunk_{chunk_index}",
                    "text": current_chunk.strip(),
                    "section_title": current_section,
                    "chunk_index": chunk_index
                })
                chunk_index += 1
                current_chunk = ""

            # 分割长段落
            for i in range(0, len(text), chunk_size - chunk_overlap):
                sub_text = text[i:i + chunk_size]
                chunks.append({
                    "chunk_id": f"chunk_{chunk_index}",
                    "text": sub_text.strip(),
                    "section_title": section_title,
                    "chunk_index": chunk_index
                })
                chunk_index += 1
        else:
            # 累积到当前 chunk
            if len(current_chunk) + len(text) + 2 <= chunk_size:
                if current_chunk:
                    current_chunk += "\n" + text
                else:
                    current_chunk = text
                if section_title and not current_section:
                    current_section = section_title
            else:
                # 当前 chunk 满了，输出并开始新的
                if current_chunk.strip():
                    chunks.append({
                        "chunk_id": f"chunk_{chunk_index}",
                        "text": current_chunk.strip(),
                        "section_title": current_section,
                        "chunk_index": chunk_index
                    })
                    chunk_index += 1

                current_chunk = text
                current_section = section_title

    # 处理最后一个 chunk
    if current_chunk.strip():
        chunks.append({
            "chunk_id": f"chunk_{chunk_index}",
            "text": current_chunk.strip(),
            "section_title": current_section,
            "chunk_index": chunk_index
        })

    return chunks


def generate_chunk_id(source_file: str, chunk_index: int) -> str:
    """生成唯一 chunk ID"""
    raw = f"{source_file}_{chunk_index}"
    return hashlib.md5(raw.encode()).hexdigest()[:16]


def get_supported_files(root_dir: str) -> List[str]:
    """递归获取目录下所有支持的文件"""
    supported_files = []
    supported_exts = ['.docx', '.pdf', '.doc', '.txt', '.xlsx', '.md']

    for root, dirs, files in os.walk(root_dir):
        # 跳过隐藏目录
        dirs[:] = [d for d in dirs if not d.startswith(".")]
        for file in files:
            if any(file.lower().endswith(ext) for ext in supported_exts):
                supported_files.append(os.path.join(root, file))

    return supported_files


def ingest_document(
    file_path: str,
    vector_store: VectorStore,
    force_doc_type: Optional[str] = None
) -> int:
    """
    摄入单个文档到向量库

    Args:
        file_path: 文件路径
        vector_store: 向量存储
        force_doc_type: 强制指定文档类型

    Returns:
        成功摄入的 chunk 数量
    """
    if force_doc_type:
        doc_type = force_doc_type
    else:
        doc_type = get_doc_type_from_path(file_path)
        if doc_type == "unknown":
            print(f"  跳过: {os.path.basename(file_path)} (无法识别文档类型)")
            return 0

    try:
        paragraphs = extract_text_from_file(file_path)
    except Exception as e:
        print(f"  读取失败: {os.path.basename(file_path)} - {e}")
        return 0

    if not paragraphs:
        print(f"  空文档: {os.path.basename(file_path)}")
        return 0

    chunks = chunk_text(paragraphs)

    # 限制每文档 chunk 数量
    if len(chunks) > MAX_CHUNKS_PER_DOC:
        chunks = chunks[:MAX_CHUNKS_PER_DOC]

    # 生成 chunk ID（基于文件名保证一致性）
    source_file = os.path.basename(file_path)
    for i, chunk in enumerate(chunks):
        chunks[i]["chunk_id"] = generate_chunk_id(source_file, i)
        chunks[i]["doc_type"] = doc_type
        chunks[i]["source_file"] = source_file

    # 批量添加到向量库
    vector_store.add_chunks(chunks)

    print(f"  [OK] {source_file}: {len(chunks)} chunks (doc_type: {doc_type})")
    return len(chunks)


def ingest_files(
    file_paths: List[str],
    collection_name: str = "all",
    force_doc_type: Optional[str] = None
) -> dict:
    """
    摄入指定文件列表到向量库

    Args:
        file_paths: 文件路径列表
        collection_name: collection 名称
        force_doc_type: 强制指定文档类型

    Returns:
        统计信息字典
    """
    vector_store = VectorStore(collection_name=collection_name)

    total_chunks = 0
    total_docs = 0

    for file_path in sorted(file_paths):
        count = ingest_document(file_path, vector_store, force_doc_type)
        if count > 0:
            total_chunks += count
            total_docs += 1

    # 文档摄入后使 BM25 缓存失效
    VectorStore.invalidate_bm25_cache()

    return {
        "total_files": len(file_paths),
        "processed_docs": total_docs,
        "total_chunks": total_chunks,
        "vector_store_count": vector_store.count()
    }


def ingest_all(
    source_dir: str,
    collection_name: str = "all",
    rebuild: bool = False
) -> dict:
    """
    摄入目录下所有文档到向量库

    Args:
        source_dir: 参考文档目录路径
        collection_name: collection 名称
        rebuild: 是否先清空现有数据

    Returns:
        统计信息字典
    """
    vector_store = VectorStore(collection_name=collection_name)

    # 先清空
    if rebuild:
        print(f"清空现有 collection: {vector_store.collection_name}")
        vector_store.clear()

    supported_files = get_supported_files(source_dir)
    print(f"发现 {len(supported_files)} 个支持的文件")

    total_chunks = 0
    total_docs = 0

    for file_path in sorted(supported_files):
        count = ingest_document(file_path, vector_store)
        if count > 0:
            total_chunks += count
            total_docs += 1

    # 文档摄入后使 BM25 缓存失效
    VectorStore.invalidate_bm25_cache()

    return {
        "total_files": len(supported_files),
        "processed_docs": total_docs,
        "total_chunks": total_chunks,
        "vector_store_count": vector_store.count()
    }


def main():
    """CLI 入口"""
    import argparse

    parser = argparse.ArgumentParser(description="RAG 文档摄入工具")
    parser.add_argument(
        "--source",
        default=str(project_root / "develop_documents"),
        help="参考文档目录路径"
    )
    parser.add_argument(
        "--files",
        nargs="+",
        help="指定文件列表摄入（空格分隔）"
    )
    parser.add_argument(
        "--collection",
        default="all",
        help="ChromaDB collection 名称"
    )
    parser.add_argument(
        "--rebuild",
        action="store_true",
        help="重建向量库（先清空）"
    )
    parser.add_argument(
        "--status",
        action="store_true",
        help="只查看向量库状态"
    )
    parser.add_argument(
        "--doc-type",
        help="强制指定文档类型"
    )

    args = parser.parse_args()

    vector_store = VectorStore(collection_name=args.collection)

    if args.status:
        count = vector_store.count()
        sources = vector_store.get_sources()
        print(f"Collection: {vector_store.collection_name}")
        print(f"总 chunk 数: {count}")
        print(f"来源文件数: {len(sources)}")
        print("\n来源文件列表:")
        for s in sources:
            print(f"  - {s}")
        return

    if args.files:
        # 摄入指定文件
        print(f"指定文件: {len(args.files)} 个")
        print(f"Collection: {args.collection}")
        if args.doc_type:
            print(f"强制文档类型: {args.doc_type}")
        print()

        result = ingest_files(
            file_paths=args.files,
            collection_name=args.collection,
            force_doc_type=args.doc_type
        )
    else:
        # 摄入整个目录
        print(f"参考文档目录: {args.source}")
        print(f"Collection: {args.collection}")
        print(f"重建模式: {args.rebuild}")
        print()

        result = ingest_all(
            source_dir=args.source,
            collection_name=args.collection,
            rebuild=args.rebuild
        )

    print()
    print("=" * 50)
    print("摄入完成!")
    print(f"  处理文档数: {result['processed_docs']}/{result['total_files']}")
    print(f"  生成 chunk 数: {result['total_chunks']}")
    print(f"  向量库总量: {result['vector_store_count']}")


if __name__ == "__main__":
    main()
