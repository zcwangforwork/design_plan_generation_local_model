"""
Template Service - Word模板加载和填充
"""

import html
import os
import re
from datetime import datetime
from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.doc_types import DOC_TYPE_LABELS, DOC_CATEGORIES
from app.services.doc_dedup import dedup_markdown

# 模板目录
TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")

# 文档类型到模板目录的映射
TEMPLATE_MAP = {
    "risk_management": "risk_management",
    "product_spec": "product_spec",
    "instruction": "instruction",
    "sop": "sop",
    "design_development_plan": "design",
    "design_input": "design",
    "design_output": "design",
    "design_review": "design",
    "design_verification": "design",
    "design_validation": "design",
    "design_change": "design",
    "design_history_file": "design"
}

# ═══════════════════════════════════════════════════════════════
# 文档前置页（首页 / 修订记录 / 目录）辅助
# ═══════════════════════════════════════════════════════════════

# 文件编号规则：{前缀}-{阶段序号}-{4位序号}
# 参考《KF-CGM-2-0004 设计输入 V7.0》：前缀=产品/公司码，阶段序号=生命周期阶段，序号=组内顺序
_CN_NUMERALS = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5,
                "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}


def _derive_doc_number(doc_type: str) -> str:
    """按阶段编号方案生成文件编号。

    阶段序号取 DOC_CATEGORIES 分组名首字的中文数字（如「二、设计输入」→ 2）；
    序号取该文档在组内 types 列表中的顺序（1 起，4 位补零）。
    未归类文档用确定性散列序号兜底（阶段 0），避免不同文档编号冲突。
    前缀可用 env DOC_NUMBER_PREFIX 覆盖（默认 KF-CGM）。
    """
    prefix = os.environ.get("DOC_NUMBER_PREFIX", "KF-CGM")
    for cat in DOC_CATEGORIES.values():
        types = cat.get("types", [])
        if doc_type in types:
            name = cat.get("name", "") or ""
            stage = _CN_NUMERALS.get(name[0], 0) if name else 0
            seq = types.index(doc_type) + 1
            return f"{prefix}-{stage}-{seq:04d}"
    seq = 1 + (sum(ord(c) for c in doc_type) % 10000)
    return f"{prefix}-0-{seq:04d}"


def _append_field(paragraph, field_code: str, placeholder: str = ""):
    """在段落末尾追加一个 Word 域（begin/instrText/separate/placeholder/end）。"""
    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(t)
    r.append(fld_end)


def _append_numpages_minus_one(paragraph):
    """追加「{ = { NUMPAGES } - 1 }」公式域，得到「总页数-1」。

    参考文档封面不编号（首页不同），故「共Y页」的 Y = 物理总页数 - 1，
    例如 15 物理页显示「共14页」、28 物理页显示「共27页」。
    """
    def _fldchar(run, ftype):
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), ftype)
        run._r.append(el)

    def _instr(run, text):
        el = OxmlElement("w:instrText")
        el.set(qn("xml:space"), "preserve")
        el.text = text
        run._r.append(el)

    def _text(run, text):
        el = OxmlElement("w:t")
        el.text = text
        run._r.append(el)

    # 外层公式域 begin + 字段码前段 " = "
    r = paragraph.add_run()
    _fldchar(r, "begin")
    _instr(r, " = ")

    # 内层 NUMPAGES 域（完整 begin/instrText/separate/result/end）
    r = paragraph.add_run()
    _fldchar(r, "begin")
    _instr(r, " NUMPAGES ")
    _fldchar(r, "separate")
    _text(r, "1")
    _fldchar(r, "end")

    # 外层公式域字段码后段 " - 1 " + separate + result + end
    r = paragraph.add_run()
    _instr(r, " - 1 ")
    _fldchar(r, "separate")
    _text(r, "1")
    _fldchar(r, "end")


def _set_update_fields_on_open(doc):
    """设置 w:updateFields，使 Word 打开文档时自动刷新 TOC/页码等域。"""
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


class TemplateService:
    """模板服务"""

    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = template_dir or TEMPLATE_DIR

    def load_template(self, doc_type: str) -> Document:
        """
        根据文档类型加载模板

        Args:
            doc_type: 文档类型

        Returns:
            Document对象
        """
        template_path = self._get_template_path(doc_type)

        if os.path.exists(template_path):
            return Document(template_path)
        else:
            # 如果模板不存在，创建空白文档
            return Document()

    def _get_template_path(self, doc_type: str) -> str:
        """获取模板文件路径"""
        template_subdir = TEMPLATE_MAP.get(doc_type, "default")
        template_path = os.path.join(
            self.template_dir,
            template_subdir,
            "template.docx"
        )
        return template_path

    def fill_template(
        self,
        doc: Document,
        content: str,
        product_name: str,
        doc_type: str
    ) -> Document:
        """
        用AI生成的内容填充模板

        Args:
            doc: Document对象
            content: AI生成的文档内容（Markdown格式）
            product_name: 产品名称
            doc_type: 文档类型

        Returns:
            填充后的Document对象
        """
        # 生成后全文档去重兜底：过滤冗余行、去除重复正文行、合并高相似度小节
        # （fill_template 是 Markdown→Word 的唯一公共转换点，覆盖所有生成路径）
        content = dedup_markdown(content)

        # 剥离"修改/精简上传文档"时模型重复输出的前置页（封面 HTML 表格、页码、目录），
        # 避免与模板自带的封面/修订记录/目录重复
        content = self._strip_preamble(content)

        # 中文字体：正文宋体、标题黑体（在写任何内容前设置，保证全局生效）
        self._set_doc_fonts(doc)

        # 前置页（首页 / 修订记录 / 目录）+ 页码页脚：每个生成的文档统一附带
        self._add_front_matter(doc, product_name, doc_type)
        self._add_page_number_footer(doc)

        # 解析Markdown内容并写入文档
        self._parse_and_fill(doc, content)
        return doc

    # ───────────────────────── 文档前置页 ─────────────────────────

    def _set_doc_fonts(self, doc):
        """设置中文字体：正文宋体（Times New Roman 数字/西文）、标题黑体。

        参考文档为中文医疗器械体系文件，正文用宋体、标题用黑体、
        标题颜色黑色（去掉 Word 默认蓝色）。字号：正文五号(10.5pt)，
        标题 H1/H2/H3 分别为 16/14/12pt。
        """
        normal = doc.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(10.5)
        self._set_east_asia_font(normal, "宋体")

        for hname, hsize in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
            st = doc.styles[hname]
            st.font.name = "Arial"
            st.font.size = Pt(hsize)
            st.font.bold = True
            st.font.color.rgb = RGBColor(0, 0, 0)
            self._set_east_asia_font(st, "黑体")

    @staticmethod
    def _set_east_asia_font(style, east_asia_name: str):
        """设置样式的东亚字体（中文字体），不影响已有西文字体设置。"""
        style.font.name = "Arial"
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), east_asia_name)

    @staticmethod
    def _set_cell(table, r: int, c: int, zh: str, en: str = None, bold: bool = False):
        """写单元格内容：中文（可加粗）+ 英文（换行显示在下方）。"""
        cell = table.cell(r, c)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(zh)
        run.bold = bold
        if en:
            run.add_break()
            p.add_run(en)
        return cell

    def _add_front_matter(self, doc, product_name: str, doc_type: str):
        """依次添加 首页 / 修订记录 / 目录 三个前置页。

        各字段：文件名称取自 DOC_TYPE_LABELS，编号由阶段编号方案生成，
        版本号/编制/审核/批准 可用 env（DOC_VERSION/DOC_MAKER/DOC_REVIEWER/DOC_APPROVER）配置。
        """
        label = DOC_TYPE_LABELS.get(doc_type, doc_type)
        doc_no = _derive_doc_number(doc_type)
        version = os.environ.get("DOC_VERSION", "V1.0")
        maker = os.environ.get("DOC_MAKER", "编制人")
        reviewer = os.environ.get("DOC_REVIEWER", "审核人")
        approver = os.environ.get("DOC_APPROVER", "批准人")
        today = datetime.now().strftime("%Y.%m.%d")

        self._add_cover(doc, label, product_name, doc_no, version,
                        maker, reviewer, approver, today)
        self._add_revision_history(doc, version, maker, reviewer, approver, today)
        self._add_toc(doc)

    def _add_front_matter_title(self, doc, text: str):
        """前置页小节标题（居中加粗，非 Heading 样式，避免进入正文目录）。"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)

    def _add_cover(self, doc, label, product_name, doc_no, version,
                   maker, reviewer, approver, today):
        """首页：双语文件信息表 + 文档标题 + 相关文档表 + 编制/审核/批准签名表。

        参照 KF- 系列受控文件封面：4 列表格，字段为
        文件类型/保密密级/文件编号/文件版本/适用范围，中英文双语。
        """
        for _ in range(2):
            doc.add_paragraph()

        # 文档信息表（4 列，10 行）
        info = doc.add_table(rows=10, cols=4, style="Table Grid")
        info.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 行0：文件类型 / 开发文档 / 保密密级 / 机密
        self._set_cell(info, 0, 0, "文件类型", "Document Type", bold=True)
        self._set_cell(info, 0, 1, "开发文档", "R&D Document")
        self._set_cell(info, 0, 2, "保密密级", "Confidentiality", bold=True)
        self._set_cell(info, 0, 3, "机密", "Confidential")
        # 行1：文件编号（合并后3列）
        self._set_cell(info, 1, 0, "文件编号", "Document No.", bold=True)
        info.cell(1, 1).merge(info.cell(1, 3))
        self._set_cell(info, 1, 1, doc_no)
        # 行2：文件版本（合并后3列）
        self._set_cell(info, 2, 0, "文件版本", "Document Version", bold=True)
        info.cell(2, 1).merge(info.cell(2, 3))
        self._set_cell(info, 2, 1, version)
        # 行3：适用范围（合并后3列）
        self._set_cell(info, 3, 0, "适用范围", "Applicable Scope", bold=True)
        info.cell(3, 1).merge(info.cell(3, 3))
        self._set_cell(info, 3, 1, product_name)
        # 行4：文档标题（合并整行，居中）
        info.cell(4, 0).merge(info.cell(4, 3))
        title_cell = self._set_cell(info, 4, 0, label, bold=True)
        for p in title_cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 行5：相关文档（合并整行，节标题）
        info.cell(5, 0).merge(info.cell(5, 3))
        self._set_cell(info, 5, 0, "相关文档", "Related Documents", bold=True)
        # 行6：相关文档表头
        self._set_cell(info, 6, 0, "文件编号", "Document No.", bold=True)
        info.cell(6, 1).merge(info.cell(6, 2))
        self._set_cell(info, 6, 1, "文件名称", "Document Name", bold=True)
        self._set_cell(info, 6, 3, "版本", "Version", bold=True)
        # 行7-9：相关文档空行（留待填写，保持表格结构完整）
        for r in (7, 8, 9):
            self._set_cell(info, r, 0, "")
            self._set_cell(info, r, 1, "")
            self._set_cell(info, r, 3, "")

        doc.add_paragraph()

        # 编制/审核/批准签名表（4 列）
        sig = doc.add_table(rows=3, cols=4, style="Table Grid")
        sig.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_rows = [
            ("编制人员", maker),
            ("审核人员", reviewer),
            ("批准人员", approver),
        ]
        for i, (role, name) in enumerate(sig_rows):
            self._set_cell(sig, i, 0, role, bold=True)
            self._set_cell(sig, i, 1, name)
            self._set_cell(sig, i, 2, "日期", bold=True)
            self._set_cell(sig, i, 3, today)

    def _add_revision_history(self, doc, version, maker, reviewer, approver, today):
        """修订记录页：双语标题 + 5列表格（版本/ECN·PCN·CR/修订内容概述/修订人/生效日期）。"""
        doc.add_page_break()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("修订记录")
        run.bold = True
        run.font.size = Pt(16)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Revision History")
        run2.bold = True
        run2.font.size = Pt(12)

        table = doc.add_table(rows=17, cols=5, style="Table Grid")
        headers = [
            ("版本", "Version"),
            ("ECN/PCN/TCN/CR", None),
            ("修订内容概述", "Description"),
            ("修订人", "Revised By"),
            ("生效日期", "Effective Date"),
        ]
        for j, (zh, en) in enumerate(headers):
            self._set_cell(table, 0, j, zh, en, bold=True)
        # 首条修订记录
        first_row = [version, "/", "新增", maker, today]
        for j, v in enumerate(first_row):
            self._set_cell(table, 1, j, v)
        # 其余空行
        for r in range(2, 17):
            for j in range(5):
                self._set_cell(table, r, j, "")

    def _add_toc(self, doc):
        """目录页：插入 Word TOC 域（Heading 1-3），打开文档时自动更新。"""
        doc.add_page_break()
        self._add_front_matter_title(doc, "目录")

        toc_para = doc.add_paragraph()
        _append_field(toc_para, 'TOC \\o "1-3" \\h \\z \\u', "（目录将在打开文档时自动更新）")
        _set_update_fields_on_open(doc)

        # 目录后分页，正文从新页开始
        doc.add_page_break()

    def _add_page_number_footer(self, doc):
        """页脚添加「第N页共Y页」页码域；首页（封面）不显示页码。

        参考文档页码格式无空格（第1页共7页），且封面不编号，
        通过「首页不同」实现：首页页脚留空，其余页显示页码。
        """
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("第")
        _append_field(p, "PAGE", "1")
        p.add_run("页共")
        _append_numpages_minus_one(p)
        p.add_run("页")

    def _add_formatted_paragraph(self, doc: Document, text: str, style: Optional[str] = None):
        """
        解析内联 Markdown 并添加为 Word 段落，支持：
        - **加粗** / __加粗__
        - *斜体* / _斜体_
        - `行内代码`
        - ~~删除线~~
        - 组合格式如 ***加粗斜体***
        """
        p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        if not text:
            return p

        # 正则匹配各种内联格式（按优先级）
        pattern = re.compile(
            r'(\*\*\*|___)(.+?)\1'           # 加粗斜体
            r'|(\*\*|__)(.+?)\3'              # 加粗
            r'|(\*|_)(.+?)\5'                 # 斜体
            r'|(`)(.+?)\7'                    # 行内代码
            r'|(~~)(.+?)\9'                   # 删除线
        )

        last_end = 0
        for match in pattern.finditer(text):
            # 添加匹配前的纯文本
            prefix = text[last_end:match.start()]
            if prefix:
                p.add_run(prefix)

            bold_italic, bi_text = match.group(1), match.group(2)
            bold, b_text = match.group(3), match.group(4)
            italic, i_text = match.group(5), match.group(6)
            code, c_text = match.group(7), match.group(8)
            strike, s_text = match.group(9), match.group(10)

            if bold_italic and bi_text:
                run = p.add_run(bi_text)
                run.bold = True
                run.italic = True
            elif bold and b_text:
                run = p.add_run(b_text)
                run.bold = True
            elif italic and i_text:
                run = p.add_run(i_text)
                run.italic = True
            elif code and c_text:
                run = p.add_run(c_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            elif strike and s_text:
                run = p.add_run(s_text)
                run.font.strike = True

            last_end = match.end()

        # 添加剩余纯文本
        remaining = text[last_end:]
        if remaining:
            p.add_run(remaining)

        return p

    def _parse_and_fill(
        self,
        doc: Document,
        content: str,
    ):
        """解析Markdown内容并填充到文档"""
        # 确保content是字符串
        if not isinstance(content, str):
            content = str(content)

        lines = content.split("\n")

        # 首页已展示文件名称与产品名称，正文直接从解析的 Markdown 内容开始
        # 解析Markdown内容
        current_heading = None
        in_code_block = False
        code_content = []
        mermaid_block = False
        mermaid_content = []
        table_buffer = []
        in_html_table = False
        html_table_buffer = []

        # 编译数字列表正则
        num_list_re = re.compile(r'^(\d+)[.)]\s+(.*)')

        for line in lines:
            stripped = line.strip()

            # 跳过标题行（第一个#开头的行是标题）
            if stripped.startswith("# ") and current_heading is None:
                continue

            # 代码块处理（含 mermaid 流程图块）
            if stripped.startswith("```"):
                self._flush_table(doc, table_buffer)
                fence_lang = stripped[3:].strip().lower()
                if in_code_block or mermaid_block:
                    # 关闭围栏
                    if mermaid_block:
                        self._add_mermaid_or_fallback(doc, mermaid_content)
                        mermaid_content = []
                        mermaid_block = False
                    else:
                        p = doc.add_paragraph()
                        p.style = "Quote"
                        p.add_run("\n".join(code_content))
                        code_content = []
                        in_code_block = False
                else:
                    # 打开围栏：mermaid 走图片渲染，其余走普通代码块
                    if fence_lang == "mermaid":
                        mermaid_block = True
                    else:
                        in_code_block = True
                continue

            if in_code_block:
                code_content.append(stripped)
                continue

            if mermaid_block:
                mermaid_content.append(stripped)
                continue

            # HTML 表格块处理（本地模型可能输出 HTML <table> 而非 Markdown 管道表）
            if in_html_table:
                html_table_buffer.append(line)
                if "</table>" in stripped.lower():
                    self._add_html_table(doc, "\n".join(html_table_buffer))
                    html_table_buffer = []
                    in_html_table = False
                continue

            if "<table" in stripped.lower():
                self._flush_table(doc, table_buffer)
                if "</table>" in stripped.lower():
                    self._add_html_table(doc, stripped)
                else:
                    in_html_table = True
                    html_table_buffer = [line]
                continue

            # 表格行处理
            if stripped.startswith("| ") or (stripped.startswith("|") and "|" in stripped[1:]):
                table_buffer.append(stripped)
                continue

            # 非表格行，先flush缓冲的表格
            self._flush_table(doc, table_buffer)

            # 标题处理
            if stripped.startswith("## "):
                current_heading = stripped[3:].strip()
                doc.add_heading(current_heading, level=1)
            elif stripped.startswith("### "):
                sub_heading = stripped[4:].strip()
                doc.add_heading(sub_heading, level=2)
            elif stripped.startswith("#### "):
                sub_heading = stripped[5:].strip()
                doc.add_heading(sub_heading, level=3)
            elif stripped == "---":
                p = doc.add_paragraph()
                p.add_run("─" * 50)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                # 无序列表 — 使用内联格式化
                self._add_formatted_paragraph(doc, stripped[2:], style="List Bullet")
            elif num_list_re.match(stripped):
                # 数字列表
                match = num_list_re.match(stripped)
                list_text = match.group(2)
                self._add_formatted_paragraph(doc, list_text, style="List Number")
            elif stripped:
                # 普通段落 — 使用内联格式化
                self._add_formatted_paragraph(doc, stripped)

        # 处理末尾可能残留的表格
        self._flush_table(doc, table_buffer)

    def _add_mermaid_or_fallback(self, doc: Document, mermaid_content: list):
        """把 mermaid 源码渲染为图片插入 Word；失败则降级为代码文本段落。

        mermaid 渲染依赖本地 node_modules + Playwright Chromium，任一不可用或
        语法错误时返回 None，这里回退为 Quote 样式的源码文本，保证文档仍能生成。
        """
        code = "\n".join(mermaid_content).strip()
        if not code:
            return

        png = None
        try:
            from app.services.mermaid_render import render_mermaid_to_png
            png = render_mermaid_to_png(code)
        except Exception:
            png = None

        if png:
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(BytesIO(png), width=Inches(5.5))
                return
            except Exception:
                pass

        # 降级：渲染失败时保留 mermaid 源码为引用样式
        p = doc.add_paragraph()
        p.style = "Quote"
        p.add_run(code)

    def _flush_table(self, doc: Document, table_buffer: list):
        """将缓冲的Markdown表格行转换为Word表格"""
        if not table_buffer:
            return

        rows = []
        has_header = False

        for row_text in table_buffer:
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            # 检测分隔行 (如 |---|---|)
            is_separator = all(
                c.replace("-", "").replace(":", "").replace(" ", "") == ""
                for c in cells
            )
            if is_separator:
                has_header = True
                continue
            rows.append(cells)

        if not rows:
            table_buffer.clear()
            return

        # 确定列数（取最宽的行）
        num_cols = max(len(row) for row in rows)

        # 填充不足的列
        for row in rows:
            while len(row) < num_cols:
                row.append("")

        # 创建Word表格
        table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")

        for i, row in enumerate(rows):
            for j in range(num_cols):
                cell = table.cell(i, j)
                cell.text = row[j]
                # 表头加粗
                if has_header and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # 表格后加空行
        doc.add_paragraph()
        table_buffer.clear()

    @staticmethod
    def _html_cell_text(inner: str) -> str:
        """把 HTML 单元格内部内容清洗为纯文本：<br> 转行、剥标签、反转义实体。"""
        inner = re.sub(r"<br\s*/?>", "\n", inner, flags=re.IGNORECASE)
        inner = re.sub(r"<[^>]+>", "", inner)
        return html.unescape(inner).strip()

    def _add_html_table(self, doc: Document, table_html: str):
        """把一段 HTML <table>...</table> 解析成 Word 表格。

        兜底本地模型不遵守"禁止输出 HTML 表格"时，把 <table> 标签正确渲染，
        而不是当成普通文本原样写进 Word。
        支持 colspan 展开（重复该单元格），忽略 rowspan。
        """
        tr_re = re.compile(r"<tr[^>]*>(.*?)</tr>", re.IGNORECASE | re.DOTALL)
        td_full_re = re.compile(r"<t[dh]\b([^>]*)>(.*?)</t[dh]>", re.IGNORECASE | re.DOTALL)

        rows = []
        for tr in tr_re.findall(table_html):
            cells = []
            for attrs, inner in td_full_re.findall(tr):
                colspan = 1
                m = re.search(r'colspan\s*=\s*["\']?(\d+)', attrs, re.IGNORECASE)
                if m:
                    try:
                        colspan = int(m.group(1))
                    except ValueError:
                        colspan = 1
                text = self._html_cell_text(inner)
                cells.extend([text] * max(colspan, 1))
            if cells:
                rows.append(cells)

        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < num_cols:
                r.append("")

        table = doc.add_table(rows=len(rows), cols=num_cols, style="Table Grid")
        for i, row in enumerate(rows):
            for j in range(num_cols):
                table.cell(i, j).text = row[j]

        # 表格后加空行
        doc.add_paragraph()

    def _strip_preamble(self, content: str) -> str:
        """剥离文档开头重复的前置页噪声（封面 HTML 表格、页码标记、目录）。

        fill_template 已自带封面/修订记录/目录；而本地模型在"修改/精简上传文档"时
        往往"原样保留"原文档的封面/修订记录/目录，导致前置页重复。这里只删除
        出现在正文第一个章标题之前的噪声，不触碰正文。
        """
        lines = content.split("\n")
        # 章标题：`# ` 标题、`第X章/第X节`、中文数字章
        chapter_re = re.compile(r"^(#+\s+|第\s*[0-9一二三四五六七八九十]+\s*[章节篇])")
        # 目录条目（以省略号+页码结尾）不算正文标题
        toc_entry_re = re.compile(r"\.{3,}\s*\d+\s*$")

        body_start = None
        for i, line in enumerate(lines):
            s = line.strip()
            if chapter_re.match(s) and not toc_entry_re.search(s):
                body_start = i
                break

        if body_start is None:
            return content

        body = lines[body_start:]
        # 移除残留的页码标记（"第X页共N页"）
        page_mark_re = re.compile(r"^第\s*\d+\s*页\s*共\s*\d+\s*页\s*$")
        body = [l for l in body if not page_mark_re.match(l.strip())]
        return "\n".join(body)

    def save_document(self, doc: Document, file_path: str):
        """保存文档到文件"""
        doc.save(file_path)

    def document_to_bytes(self, doc: Document) -> bytes:
        """将文档转换为字节流"""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        # 确保返回的是字节而不是字符串
        return buffer.getvalue()
